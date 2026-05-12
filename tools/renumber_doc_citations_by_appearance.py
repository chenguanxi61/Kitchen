from pathlib import Path
import re
import shutil

from docx import Document


DOC = Path(r"E:\UnityLearn\Kitchen\output\doc\基于Unity的多人协作烹饪游戏设计与实现.docx")
BACKUP = DOC.with_name("基于Unity的多人协作烹饪游戏设计与实现_引用顺序调整前备份.docx")


def find_reference_start(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "参考文献":
            return i
    raise RuntimeError("未找到参考文献标题")


def replace_paragraph_text(paragraph, text):
    paragraph.text = text


def map_citation_group(match, mapping):
    group = match.group(0)
    old_numbers = [int(n) for n in re.findall(r"\[(\d+)\]", group)]
    new_numbers = [mapping[n] for n in old_numbers if n in mapping]
    new_numbers = sorted(dict.fromkeys(new_numbers))
    return "".join(f"[{n}]" for n in new_numbers)


if not BACKUP.exists():
    shutil.copyfile(DOC, BACKUP)

doc = Document(DOC)
ref_start = find_reference_start(doc)
citation_pattern = re.compile(r"\[(\d+)\]")
citation_group_pattern = re.compile(r"(?:\[\d+\])+")

first_appearance = []
for paragraph in doc.paragraphs[:ref_start]:
    for number in citation_pattern.findall(paragraph.text):
        old = int(number)
        if old not in first_appearance:
            first_appearance.append(old)

mapping = {old: new for new, old in enumerate(first_appearance, start=1)}

for paragraph in doc.paragraphs[:ref_start]:
    text = paragraph.text
    if not citation_pattern.search(text):
        continue
    new_text = citation_group_pattern.sub(lambda m: map_citation_group(m, mapping), text)
    replace_paragraph_text(paragraph, new_text)

reference_paragraphs = []
for i in range(ref_start + 1, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if not text:
        continue
    if text.startswith("致"):
        break
    reference_paragraphs.append(doc.paragraphs[i])

reference_text_by_old_number = {
    old_number: reference_paragraphs[old_number - 1].text
    for old_number in range(1, len(reference_paragraphs) + 1)
}

for new_index, old_number in enumerate(first_appearance):
    reference_paragraphs[new_index].text = reference_text_by_old_number[old_number]

doc.save(DOC)

print("mapping old->new:", mapping)
print(DOC)
print(BACKUP)
