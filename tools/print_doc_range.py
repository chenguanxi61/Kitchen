from pathlib import Path
import sys

from docx import Document


doc_path = Path(sys.argv[1])
start = int(sys.argv[2])
end = int(sys.argv[3])

doc = Document(doc_path)
for i in range(start, min(end, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    if text:
        print(f"{i}: {text}")
