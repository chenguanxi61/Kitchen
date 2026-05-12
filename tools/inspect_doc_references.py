from pathlib import Path

from docx import Document


DOC = Path(r"E:\UnityLearn\Kitchen\output\doc\基于Unity的多人协作烹饪游戏设计与实现.docx")


def para_texts(document):
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


doc = Document(DOC)
texts = para_texts(doc)

print("HEADINGS")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ""
    if "Heading" in style or text[:3] in {"第1章", "第2章", "第3章", "第4章", "第5章", "第6章", "参考文"} or text.startswith(("1.", "2.", "3.", "4.", "5.", "6.")):
        print(f"{i}: [{style}] {text[:100]}")

print("\nFIRST CHAPTER")
in_first = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    if text.startswith(("第1章", "第 1 章", "1 ")):
        in_first = True
    elif in_first and text.startswith(("第2章", "第 2 章", "2 ")):
        break
    if in_first:
        print(f"{i}: {text}")

print("\nREFERENCES")
in_ref = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text in {"参考文献", "参 考 文 献"}:
        in_ref = True
    if in_ref:
        print(f"{i}: {text}")
