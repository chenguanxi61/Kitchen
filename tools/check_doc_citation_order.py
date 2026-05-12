from pathlib import Path
import re

from docx import Document


DOC = Path(r"E:\UnityLearn\Kitchen\output\doc\基于Unity的多人协作烹饪游戏设计与实现.docx")

doc = Document(DOC)
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "参考文献":
        ref_start = i
        break

pattern = re.compile(r"\[(\d+)\]")
seen = []
locations = []
for i, p in enumerate(doc.paragraphs[:ref_start]):
    text = p.text
    for m in pattern.finditer(text):
        num = int(m.group(1))
        if num not in seen:
            seen.append(num)
            locations.append((num, i, text.strip()[:140]))

print("FIRST APPEARANCE ORDER")
print(seen)
for num, i, text in locations:
    print(f"[{num}] first at paragraph {i}: {text}")

print("\nREFERENCES")
for i, p in enumerate(doc.paragraphs[ref_start + 1 :], start=ref_start + 1):
    text = p.text.strip()
    if not text:
        continue
    if text.startswith("致"):
        break
    print(f"{i}: {text}")
